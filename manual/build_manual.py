"""Build the archived v0.5.3 DOCX manual from its original screenshots.

The current manuals are USER_GUIDE.md and USER_GUIDE_zh-TW.md.  This builder
is intentionally versioned as a legacy source because its screenshots and
instructions describe the v0.5.3 interface.
"""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RAW = ROOT / "manual" / "assets" / "raw"
ANN = ROOT / "manual" / "assets" / "annotated"
LEGACY_VERSION = "0.5.3"
OUT = (
    ROOT
    / "manual"
    / f"S2P-XInput-Lite-legacy-v{LEGACY_VERSION}-繁體中文使用手冊.docx"
)
ANN.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
NAVY = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F2F7FB"
ORANGE = "D97706"
RED = "C62828"
GREEN = "178A38"
GRAY = "5E6670"
LIGHT_GRAY = "F3F4F6"
BLACK = "111827"
WHITE = "FFFFFF"
FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft JhengHei"
NEXT_PAGE_BREAK = False


def find_font(size=24, bold=True):
    candidates = [
        Path(r"C:\Windows\Fonts\msjhbd.ttc" if bold else r"C:\Windows\Fonts\msjh.ttc"),
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def annotate(name, points):
    src = RAW / f"{name}.png"
    out = ANN / f"{name}.png"
    image = Image.open(src).convert("RGB")
    draw = ImageDraw.Draw(image)
    font = find_font(22, True)
    for number, x, y in points:
        radius = 17
        draw.ellipse(
            (x - radius, y - radius, x + radius, y + radius),
            fill=(198, 40, 40),
            outline=(255, 255, 255),
            width=3,
        )
        label = str(number)
        box = draw.textbbox((0, 0), label, font=font)
        tw, th = box[2] - box[0], box[3] - box[1]
        draw.text((x - tw / 2, y - th / 2 - 2), label, font=font, fill=(255, 255, 255))
    image.save(out, quality=95)
    return out


FIGURES = {
    "main": annotate(
        "01-main-buttons",
        [(1, 183, 195), (2, 186, 436), (3, 187, 692), (4, 522, 64),
         (5, 534, 438), (6, 535, 826), (7, 177, 885), (8, 532, 924)],
    ),
    "stick_map": annotate(
        "02-stick-map",
        [(1, 548, 204), (2, 426, 335), (3, 598, 335), (4, 430, 753)],
    ),
    "layers": annotate(
        "03-layers",
        [(1, 383, 116), (2, 603, 115), (3, 472, 141), (4, 535, 791)],
    ),
    "rumble_game": annotate(
        "04-advanced-rumble-game",
        [(1, 521, 116), (2, 418, 381), (3, 562, 509), (4, 515, 686)],
    ),
    "rumble_audio": annotate(
        "05-advanced-rumble-audio",
        [(1, 548, 116), (2, 522, 223), (3, 523, 410), (4, 562, 509)],
    ),
    "gyro": annotate(
        "06-gyro-map",
        [(1, 525, 121), (2, 520, 188), (3, 522, 278), (4, 526, 480), (5, 584, 684)],
    ),
    "curve": annotate(
        "07-stick-curve-zoom",
        [(1, 229, 581), (2, 400, 393), (3, 571, 219), (4, 313, 851)],
    ),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=11, bold=False, color=BLACK, italic=False):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(p, text, **kwargs):
    run = p.add_run(text)
    set_font(run, **kwargs)
    return run


def format_paragraph(p, before=0, after=6, line=1.25, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep


def add_heading(doc, text, level=1):
    global NEXT_PAGE_BREAK
    p = doc.add_paragraph()
    sizes = {1: 16, 2: 13, 3: 12}
    colors = {1: BLUE, 2: BLUE, 3: NAVY}
    before = {1: 18, 2: 14, 3: 10}
    after = {1: 10, 2: 7, 3: 5}
    format_paragraph(p, before[level], after[level], 1.0, True)
    if NEXT_PAGE_BREAK:
        p.paragraph_format.page_break_before = True
        NEXT_PAGE_BREAK = False
    add_text(p, text, size=sizes[level], bold=True, color=colors[level])
    return p


def add_body(doc, text, bold_prefix=None, color=BLACK, after=6):
    p = doc.add_paragraph()
    format_paragraph(p, 0, after, 1.25)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True, color=color)
        add_text(p, text[len(bold_prefix):], color=color)
    else:
        add_text(p, text, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.375)
        pf.first_line_indent = Inches(-0.188)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.25
        add_text(p, "• ", bold=True, color=BLUE)
        add_text(p, item)


def add_steps(doc, items):
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.375)
        pf.first_line_indent = Inches(-0.188)
        pf.space_after = Pt(5)
        pf.line_spacing = 1.25
        add_text(p, f"{idx}. ", bold=True, color=BLUE)
        add_text(p, item)


def add_note(doc, title, text, kind="info"):
    colors = {
        "info": (PALE_BLUE, BLUE),
        "warn": ("FFF7E8", ORANGE),
        "danger": ("FDECEC", RED),
        "ok": ("ECF7EF", GREEN),
    }
    fill, accent = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 100, 160, 100, 160)
    p = cell.paragraphs[0]
    format_paragraph(p, 0, 2, 1.2)
    add_text(p, title + "　", bold=True, color=accent)
    add_text(p, text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc, path, caption, width=4.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(cap, 0, 8, 1.0)
    add_text(cap, caption, size=9, color=GRAY, italic=True)


def add_callouts(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.55, 5.95]
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    hdr = table.rows[0].cells
    hdr[0].text = "編號"
    hdr[1].text = "說明"
    set_repeat_table_header(table.rows[0])
    for cell in hdr:
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            set_font(run, size=10, bold=True, color=NAVY)
    for num, text in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(cells[0].paragraphs[0], str(num), bold=True, color=RED)
        add_text(cells[1].paragraphs[0], text, size=10)
        for cell in cells:
            set_cell_margins(cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_break(doc):
    global NEXT_PAGE_BREAK
    NEXT_PAGE_BREAK = True


def add_contents(doc):
    add_heading(doc, "目錄", 1)
    rows = [
        ("1", "開始使用", "安裝需求、第一次啟動、三種連線方式"),
        ("2", "主介面與方案", "狀態列、儲存、建立、匯入與管理方案"),
        ("3", "搖桿", "曲線、死區、防抖、直接輸入與還原"),
        ("4", "按鍵與方向映射", "Xbox 按鍵、四向映射與觸發條件"),
        ("5", "Mapping Layers", "建立、啟用與切換額外映射層"),
        ("6", "陀螺儀", "啟用方式、目標、控制模式、校正與感度"),
        ("7", "震動", "LF/HF、最大振幅、Game／Audio／Mix 與六頻段"),
        ("8", "HidHide 與 ESP32", "避免雙重輸入、韌體刷寫與配對"),
        ("9", "操作技巧與疑難排解", "快速輸入、還原、常見症狀"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.45, 1.65, 4.4]
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    for idx, label in enumerate(("章", "主題", "內容")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        add_text(cell.paragraphs[0], label, bold=True, color=NAVY)
        set_cell_margins(cell)
    set_repeat_table_header(table.rows[0])
    for chapter, topic, desc in rows:
        cells = table.add_row().cells
        for i, width in enumerate(widths):
            cells[i].width = Inches(width)
            set_cell_margins(cells[i])
        add_text(cells[0].paragraphs[0], chapter, bold=True, color=BLUE)
        add_text(cells[1].paragraphs[0], topic, bold=True)
        add_text(cells[2].paragraphs[0], desc, size=10)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_paragraph(header, 0, 0, 1.0)
    add_text(header, "S2P-XInput-Lite｜繁體中文使用手冊", size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(footer, 0, 0, 1.0)
    add_text(
        footer,
        f"legacy v{LEGACY_VERSION}　•　2026-07-23",
        size=8.5,
        color=GRAY,
    )

    # editorial_cover pattern
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(p, 76, 12, 1.0)
    add_text(p, "圖文操作指南", size=10.5, bold=True, color=ORANGE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(title, 0, 7, 1.0)
    add_text(title, "S2P-XInput-Lite", size=30, bold=True, color=NAVY)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(sub, 0, 20, 1.0)
    add_text(
        sub,
        f"legacy v{LEGACY_VERSION} 繁體中文使用手冊",
        size=16,
        color=BLUE,
    )
    banner = ROOT / "image" / "S2P-XInput-Lite-banner.jpg"
    if banner.exists():
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(banner), width=Inches(5.5))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(p, 22, 3, 1.0)
    add_text(p, "安裝・連線・映射・陀螺儀・震動・ESP32", size=11, bold=True, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(p, 0, 0, 1.0)
    add_text(
        p,
        f"適用版本：v{LEGACY_VERSION}",
        size=10,
        italic=True,
        color=GRAY,
    )

    add_page_break(doc)
    add_heading(doc, "使用這份手冊", 1)
    add_body(doc, f"本手冊使用 v{LEGACY_VERSION} 實際介面截圖，紅色圓形編號會對應到圖下方的說明。截圖採英文介面；按下主視窗左下角「中 / En」即可切換語言，功能位置不變。")
    add_note(doc, "重要", "修改介面上的數值後，必須按「儲存/套用（Save/Apply）」才會寫入目前方案並套用到連線。只切換方案或關閉視窗，不代表已儲存。", "warn")
    add_note(doc, "設定安全", "建議先建立自己的方案，再逐項調整。System Default 是唯讀基準；若結果不理想，可將單一參數還原到上次儲存值或系統預設。", "info")
    add_contents(doc)

    add_page_break(doc)
    add_heading(doc, "1　開始使用", 1)
    add_heading(doc, "1.1　系統需求", 2)
    add_bullets(doc, [
        "Windows 10 或 Windows 11。",
        "相容控制器與 USB-C 線材；亦可使用 ESP32-S3 橋接或 Windows BLE。",
        "ViGEmBus 1.22.0，用來建立虛擬 Xbox 360 控制器。",
        "USB 有線模式建議安裝 HidHide，避免遊戲同時收到實體 HID 與虛擬 Xbox 輸入。",
    ])
    add_heading(doc, "1.2　第一次啟動", 2)
    add_steps(doc, [
        "將發佈壓縮檔完整解壓到新的資料夾，不要直接在壓縮檔內執行。",
        "第一次使用時，執行 driver 資料夾中的 Install-ViGEmBus.bat 安裝虛擬控制器驅動。",
        "連接控制器或 ESP32-S3，然後執行 S2P-XInput-Lite.exe。",
        "確認視窗底部 ViGEm 顯示綠色 Ready。依連線方式確認 Pad／ESP32 狀態。",
        "選擇方案，調整設定後按「儲存/套用」。",
    ])
    add_note(doc, "更新版本", "由舊版更新時，請解壓到新資料夾。若要保留個人設定，可在程式關閉時複製舊版 src/config.ini 到新版。", "info")
    add_heading(doc, "1.3　三種連線方式", 2)
    add_callouts(doc, [
        ("USB", "控制器直接以 USB-C 連接電腦。延遲低；建議搭配 HidHide。"),
        ("ESP32", "控制器以 BLE 連接 ESP32-S3，再由 ESP32 透過 USB 傳給電腦。"),
        ("Windows BLE", "控制器直接與 Windows 配對。方便，但更新率與穩定度可能受系統藍牙環境影響。"),
    ])

    add_page_break(doc)
    add_heading(doc, "2　主介面與方案", 1)
    add_figure(doc, FIGURES["main"], "圖 2-1　主介面總覽（紅色編號對應下表）", 4.55)
    add_callouts(doc, [
        (1, "搖桿曲線：拖曳控制點調整輸入與輸出的對應。Left／Right 分別設定左右搖桿。"),
        (2, "輸出與防抖：Shape 調整輸出形狀；Stab. 抑制曲線放大區域的小幅抖動。"),
        (3, "震動基本設定：LF/HF 強度、曲線、混合、頻率與最大振幅。"),
        (4, "功能頁籤：Buttons、Stick Map、Layers、Advanced Rumble、Gyro Map。"),
        (5, "目前頁面內容：圖中為 Switch 2 Pro 按鍵對應 Xbox 按鍵。"),
        (6, "狀態列：ViGEm、WASAPI、HidHide 與控制器／ESP32 連線狀態。"),
        (7, "方案選單與儲存工具：Save/Apply、Save New、Import Profile、Profile Mgr.。"),
        (8, "全域工具：切換語言、還原、校正、刷韌體、重啟連線與 Pin。"),
    ])
    add_heading(doc, "2.1　方案操作", 2)
    add_bullets(doc, [
        "切換方案：從 Profile 下拉選單選擇已儲存方案。",
        "儲存/套用：覆寫目前可寫入的方案，並立即套用至連線。",
        "另存新方案：保留目前方案，使用目前畫面數值建立新方案。",
        "匯入方案：載入外部 .ini 方案；名稱衝突時依視窗提示處理。",
        "管理方案：重新命名或刪除個人方案。System Default 不可覆寫。",
    ])

    add_page_break(doc)
    add_heading(doc, "3　搖桿設定", 1)
    add_figure(doc, FIGURES["curve"], "圖 3-1　放大的搖桿曲線編輯器", 5.2)
    add_callouts(doc, [
        (1, "25% 控制點：調整小幅推桿的輸出反應。"),
        (2, "50% 控制點：調整中段反應。"),
        (3, "75% 控制點：調整接近外圈時的反應。"),
        (4, "Shape 與 Stab.：控制曲線形狀與防抖補償。"),
    ])
    add_heading(doc, "3.1　曲線與死區", 2)
    add_bullets(doc, [
        "拖曳藍色控制點可改變曲線。左搖桿為藍色，右搖桿為紅色。",
        "雙擊控制點可還原該點預設；右鍵控制點可直接輸入 X、Y 座標。",
        "CTR DZ 是中心死區；OUT DZ 是外圈死區。數值過高會縮短可用行程。",
        "Lin. 使用線性連接；Smo. 使用平滑曲線。",
        "Zoom 開啟大型編輯視窗，適合精細調整。",
    ])
    add_note(doc, "建議順序", "先完成搖桿校正，再調中心死區，最後調整曲線。若先用曲線補償硬體偏移，容易造成左右方向手感不一致。", "ok")
    add_heading(doc, "3.2　校正", 2)
    add_steps(doc, [
        "按主視窗下方「校正搖桿（Calibrate）」。",
        "依命令列提示讓搖桿回到中心，再沿完整外圈緩慢旋轉。",
        "完成後回到主介面測試；校正資料依控制器保存。",
    ])

    add_page_break(doc)
    add_heading(doc, "4　按鍵與方向映射", 1)
    add_heading(doc, "4.1　Buttons", 2)
    add_body(doc, "Buttons 頁將控制器的實體按鍵映射為 Xbox 360 輸出。左欄是實體按鍵，右側下拉選單是輸出。Reset 會還原本頁映射；仍需按 Save/Apply 才會保存。")
    add_bullets(doc, [
        "一般配置常見 A/B 與 X/Y 對調，以符合 Xbox 的按鍵位置。",
        "CAPT、C、GR、GL 等按鍵可設為任一 Xbox 按鍵或 NONE。",
        "若同一輸出被多個實體按鍵使用，兩者都會觸發該輸出。",
    ])
    add_figure(doc, FIGURES["stick_map"], "圖 4-1　Stick Map 方向映射", 4.55)
    add_callouts(doc, [
        (1, "Mode：選擇 4WAY 等方向判定方式。"),
        (2, "方向缺口與圓形圖：顯示目前方向判定區域。"),
        (3, "DZ／Trig／Rel：設定方向死區、觸發門檻與釋放門檻。"),
        (4, "左右搖桿可分別設定；Reset 只還原對應搖桿的方向映射。"),
    ])
    add_note(doc, "避免抖動", "Trig 應高於 Rel，形成遲滯區，避免搖桿停在方向邊界時快速重複觸發。", "info")

    add_page_break(doc)
    add_heading(doc, "5　Mapping Layers", 1)
    add_figure(doc, FIGURES["layers"], "圖 5-1　額外映射層管理", 4.55)
    add_callouts(doc, [
        (1, "勾選方塊可啟用／停用該層。"),
        (2, "指定啟動按鍵與 Toggle／Hold 等行為。"),
        (3, "Edit、Rename、Delete 用於編輯、重新命名與刪除。"),
        (4, "+ Add 建立新層；Import 匯入；Layer Mgr. 管理層檔案。"),
    ])
    add_heading(doc, "5.1　建立額外映射", 2)
    add_steps(doc, [
        "按「+ Add」建立映射層並輸入容易辨識的名稱。",
        "按 Edit 設定該層啟用時的按鍵、鍵盤或滑鼠映射。",
        "選擇啟動按鍵，以及按住（Hold）或切換（Toggle）方式。",
        "勾選層左側方塊，最後按 Save/Apply。",
    ])
    add_note(doc, "編輯尚未儲存", "Mapping Layer 的變更會先保留在記憶體中。只有 Save/Apply 成功後，相關層檔案與方案狀態才視為正式保存。", "warn")

    add_page_break(doc)
    add_heading(doc, "6　陀螺儀映射", 1)
    add_figure(doc, FIGURES["gyro"], "圖 6-1　Gyro Map 設定", 4.55)
    add_callouts(doc, [
        (1, "Activation：指定啟用按鍵，並選擇 Off、Hold 或 Toggle。"),
        (2, "Target：輸出到左搖桿、右搖桿或滑鼠；可反轉 X／Y。"),
        (3, "Control Mode：Aim（Center）適合瞄準；Wheel（Tilt）適合傾斜操控。"),
        (4, "Gyro Response：感度、X/Y 比例、死區、反死區與平滑。"),
        (5, "Stability 與校正工具：抑制突變、適應輸入並執行 Sensor Cal。"),
    ])
    add_heading(doc, "6.1　建議設定流程", 2)
    add_steps(doc, [
        "將控制器放在穩定、無震動的平面，執行 Sensor Cal。",
        "先選 Target，再決定 Hold 或 Toggle 啟用方式。",
        "從較低感度開始測試，調整 X/Y Ratio 修正水平與垂直速度差。",
        "只有中心容易漂移時才增加 DZ；若小動作無法輸出，再少量增加 Anti DZ。",
        "最後調 Smooth ms。過高會穩定但增加黏滯感。",
    ])
    add_note(doc, "校正環境", "進行磁力計步驟時，請遠離喇叭、磁鐵、大型金屬物與通電設備，並依提示讓控制器在多個方向做完整 3D 八字翻轉。", "warn")

    add_page_break(doc)
    add_heading(doc, "7　震動設定", 1)
    add_heading(doc, "7.1　基本震動", 2)
    add_bullets(doc, [
        "LF Strength／HF Strength：調整低頻與高頻路徑強度。",
        "LF Curve／HF Curve：改變小訊號到大訊號的反應曲線。",
        "HF → LF Mix／LF → HF Mix：將部分訊號交叉送往另一條路徑。",
        "LF Frequency／HF Frequency：設定兩個震動頻率命令。",
        "Max Amp：限制最終振幅；協議欄位範圍為 0–1023，預設 800。",
    ])
    add_note(doc, "振幅與異音", "實測上，震動強度越高，線性馬達在快速變化時越可能出現輕微撞擊聲。程式預設 Max Amp 800，約為欄位上限的 78%（可概稱約 80%）；這是通用折衷值，不是硬體安全認證上限。若手把出現明顯異音，請降低 Max Amp 或個別 LF/HF 強度。", "danger")
    add_heading(doc, "7.2　Game／Audio／Mix", 2)
    add_figure(doc, FIGURES["rumble_game"], "圖 7-1　進階震動（Game 模式）", 4.55)
    add_callouts(doc, [
        (1, "Source：Game 只使用遊戲原生震動；Audio 將系統輸出轉為震動；Mix 柔和混合兩者。"),
        (2, "六頻段曲線：拖曳控制點調整各音訊頻段的增益。"),
        (3, "LF/HF Balance：控制中間頻段偏向低頻或高頻路徑。"),
        (4, "Final Output：Tail 與 Decay 控制震動尾端與衰減。"),
    ])
    add_page_break(doc)
    add_heading(doc, "7.3　音訊震動與六頻段", 2)
    add_figure(doc, FIGURES["rumble_audio"], "圖 7-2　Audio 模式會啟用 Audio Response", 4.55)
    add_callouts(doc, [
        (1, "Audio：將 Windows 預設輸出裝置的音訊轉為震動。"),
        (2, "Lvl、Gate、Atk、Rel：控制敏感度、噪音門檻、起振與釋放速度。"),
        (3, "六頻段：Low、L-Mid、Mid、H-Mid、High、Ultra。"),
        (4, "LF/HF Balance：-1 偏 LF、0 平衡、+1 偏 HF；建議先從 -0.15～+0.15 測試。"),
    ])
    add_callouts(doc, [
        ("Low", "20–120 Hz：深沉撞擊與超低頻。"),
        ("L-Mid", "120–300 Hz：低頻厚度。"),
        ("Mid", "300–700 Hz：中低頻細節。"),
        ("H-Mid", "700–2000 Hz：中高頻細節。"),
        ("High", "2000–4000 Hz：摩擦與人聲邊緣。"),
        ("Ultra", "4000–8000 Hz：尖銳提示與超高頻細節。"),
    ])
    add_note(doc, "高頻噪音", "若 Audio 或 Mix 模式出現持續高頻聲，先降低 High／Ultra 增益、將 LF/HF Balance 稍微向 LF 移動，並提高 Gate。若所有聲音一開始就過強，再降低 Lvl。", "warn")

    add_page_break(doc)
    add_heading(doc, "8　HidHide 與 ESP32", 1)
    add_heading(doc, "8.1　HidHide", 2)
    add_body(doc, "USB 有線模式下，遊戲可能同時看到實體控制器與虛擬 Xbox 控制器，造成雙重輸入。HidHide 可隱藏實體 HID，同時允許 S2P-XInput-Lite 繼續讀取。")
    add_bullets(doc, [
        "HidHide：正常／Ready：已安裝且設定完成。",
        "HidHide：缺少／Missing：未安裝。點擊狀態文字可開啟官方下載頁。",
        "HidHide：關閉／設定：已安裝但尚未完成隱藏設定。點擊狀態可重新執行設定。",
        "選擇不再提醒後，程式不會每次啟動重複詢問；仍可從底部狀態手動開啟。",
    ])
    add_note(doc, "設定失敗", "先關閉 HidHide Configuration Client，再回到 S2P-XInput-Lite 重試。其他已在 HidHide 清單中的裝置也可能受到全域隱藏影響，請先確認清單。", "warn")
    add_heading(doc, "8.2　刷入 ESP32-S3 韌體", 2)
    add_steps(doc, [
        "按主視窗底部「刷入相容韌體（Flash FW）」。",
        "將 ESP32-S3 的 OTG 接口接到電腦。",
        "按住 BOOT，按一下 RESET／EN，先放開 RESET／EN，再放開 BOOT。",
        "程式會自動偵測新的 COM Port 並刷入韌體；過程中不要拔除裝置。",
        "完成後按 RESET／EN，或重新插拔 ESP32-S3，再重新啟動程式。",
    ])
    add_heading(doc, "8.3　控制器與 ESP32 配對", 2)
    add_steps(doc, [
        "先讓 ESP32-S3 與電腦連接並啟動 S2P-XInput-Lite。",
        "讓控制器進入配對模式；需要新配對時按住控制器 SYNC。",
        "程式偵測到 SYNC 配對連線後，會將 ESP32 的配對資料寫入控制器。",
        "確認底部 Pad／ESP32 狀態變為已連線，再進入遊戲測試。",
    ])

    add_page_break(doc)
    add_heading(doc, "9　快速操作技巧", 1)
    add_heading(doc, "9.1　滑桿與數值", 2)
    add_bullets(doc, [
        "點擊滑桿目前數值，可開啟固定寬度的參數輸入視窗；視窗同時顯示範圍、步進與說明。",
        "在參數文字上按住並向右拖曳可增加，向左拖曳可減少；每次依該參數步進變化。",
        "右鍵參數可選擇「還原至上次儲存數值」或「還原至系統預設」。",
        "搖桿曲線控制點可雙擊還原；右鍵可直接輸入 X、Y。",
        "六頻段說明文字移入滑鼠後，游標旁會出現問號提示，顯示完整頻段與 LF/HF Balance 說明。",
    ])
    add_heading(doc, "9.2　Pin 與 Restart", 2)
    add_bullets(doc, [
        "Restart：重新啟動連接程序，使已儲存設定與連線狀態重新載入。",
        "Pin：將目前偵測到的控制器作為優先裝置；成功時會有提示震動。",
        "若 Pin 只震一下或無法完成，確認控制器已連線且底部狀態不是 Searching。",
    ])
    add_note(doc, "未儲存變更", "使用 Restart 前先按 Save/Apply。否則畫面上的未儲存修改可能不會套用。", "warn")

    add_page_break(doc)
    add_heading(doc, "10　疑難排解", 1)
    add_callouts(doc, [
        ("ViGEm 不是 Ready", "重新安裝 ViGEmBus，重新開機後再啟動程式。"),
        ("遊戲收到兩組輸入", "USB 模式請安裝並設定 HidHide；也要檢查 Steam Input 是否重複轉換。"),
        ("找不到控制器", "確認線材可傳輸資料、重新插拔，或清除 Windows 藍牙配對後重新配對。"),
        ("ESP32 一直 Searching", "確認韌體版本、USB OTG 接口、COM 裝置與控制器配對狀態。"),
        ("音訊震動無反應", "確認 WASAPI 為 Ready、Windows 預設輸出裝置正確，並選擇 Audio 或 Mix。"),
        ("安靜處仍有震動", "提高 Gate，降低 Lvl，並降低 High／Ultra 頻段。"),
        ("震動有撞擊聲", "降低 Max Amp、LF/HF Strength，或降低容易觸發異音的頻段增益。"),
        ("陀螺儀漂移", "重新 Sensor Cal；降低環境震動與磁性干擾，必要時增加少量 DZ。"),
        ("方案無法覆寫", "System Default 為唯讀；使用 Save New 建立個人方案。"),
        ("設定重啟後消失", "確認修改後已按 Save/Apply，且程式所在資料夾可寫入。"),
    ])
    add_heading(doc, "10.1　建議的安全回復流程", 2)
    add_steps(doc, [
        "先將單一可疑參數還原至上次儲存值。",
        "若問題仍在，切換至 System Default 比較。",
        "只有確定要全面重設時才按 Defaults；搖桿校正與已儲存方案會保留，但 Mapping Layers 會停用，程式的 HidHide 隱藏設定也會取消。",
        "重新套用後按 Restart，再測試遊戲輸入與震動。",
    ])

    add_page_break(doc)
    add_heading(doc, "附錄　預設調整原則", 1)
    add_callouts(doc, [
        ("搖桿", "校正 → 中心死區 → 外圈死區 → 曲線 → 防抖。一次只改一類參數。"),
        ("陀螺儀", "校正 → 目標與啟用方式 → 感度 → 比例 → 死區 → 平滑。"),
        ("遊戲震動", "先用 Game 與 Max Amp 800 測試；有異音再降低上限或強度。"),
        ("音訊震動", "先調 Gate 與 Lvl，再調六頻段，最後才調 LF/HF Balance 與尾端。"),
        ("方案", "保留一個已驗證的穩定方案，實驗性設定另存新方案。"),
    ])
    add_note(doc, "版本適用性", f"本手冊以 S2P-XInput-Lite v{LEGACY_VERSION} 為準。後續版本若調整介面名稱、參數範圍或連線流程，應以新版程式內的問號說明與發佈說明為準。", "info")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
